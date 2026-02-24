#!/usr/bin/env python3
"""
Generate three dummy .docx test files for the patent claim verifier:
  test_id.docx          — Invention Disclosure
  test_additional.docx  — Additional supporting information
  test_claim.docx       — Patent Claims  (contains 3 Word comments used as questions)

Run:  python create_test_docs.py
"""
import os
import zipfile
from pathlib import Path

from docx import Document

OUT = Path(__file__).parent
INVENTION = "Smart Water Bottle with IoT Temperature & Hydration Monitoring"


# ── tiny helpers ──────────────────────────────────────────────────────────────

def h(doc, text, level=1):
    doc.add_heading(text, level=level)

def p(doc, text):
    doc.add_paragraph(text)

def b(doc, text):
    doc.add_paragraph(text, style="List Bullet")


# ── 1. Invention Disclosure ───────────────────────────────────────────────────

def make_id():
    doc = Document()
    h(doc, "Invention Disclosure", 0)
    p(doc, f"Title:     {INVENTION}")
    p(doc, "Date:      2024-01-10   |   Inventors: Alice Chen, Bob Kumar")

    h(doc, "1. Field of the Invention")
    p(doc, (
        "This invention relates to consumer hydration products — specifically a portable "
        "water bottle with an embedded IoT sensor module capable of monitoring liquid "
        "temperature and daily hydration intake simultaneously and in real time."
    ))

    h(doc, "2. Background")
    p(doc, (
        "Dehydration causes measurable cognitive and physical performance decline. "
        "Existing smart bottles measure either weight (volume) or temperature, not both. "
        "None provide continuous wireless synchronisation with mobile health applications "
        "alongside active hydration-deficit alerts."
    ))

    h(doc, "3. Summary of the Invention")
    p(doc, "Key components of the invention:")
    for item in [
        "Double-walled stainless-steel bottle body (400 ml – 1 L variants).",
        "NTC thermistor sensor accurate to ±0.5 °C, seated in the base cap.",
        "Strain-gauge load cell measuring liquid mass to ±2 ml accuracy.",
        "Bluetooth Low Energy (BLE) 5.0 microcontroller — Nordic nRF52840.",
        "Companion mobile application (iOS / Android) with real-time dashboard.",
        "Rechargeable 80 mAh Li-Po battery providing ≥ 14 days per charge.",
        "Wireless OTA firmware update capability via the mobile app.",
    ]:
        b(doc, item)

    h(doc, "4. Detailed Description")
    p(doc, (
        "The sensor module occupies a waterproof compartment inside the bottle base "
        "(IP67-rated, TPU-gasket sealed). The NTC thermistor connects to a 12-bit ADC "
        "on the microcontroller; temperature is sampled every 10 s and pushed over BLE "
        "GATT notify packets. The load cell is sampled at 100 ms with a 5-sample moving "
        "average to suppress vibration artefacts. "
        "The mobile app aggregates readings, computes daily hydration deficit against a "
        "body-weight-derived target (35 ml/kg), and issues push-notification reminders. "
        "All data is persisted on-device; optional cloud sync uses a REST API over TLS 1.3."
    ))

    h(doc, "5. Claims Sought")
    for item in [
        "Broad apparatus claim: portable container with combined temp + volume sensing.",
        "Method claim: computing hydration deficit from body-weight-based daily target.",
        "System claim: bottle + mobile app + cloud service as an integrated product.",
    ]:
        b(doc, item)

    h(doc, "6. Prior Art Known to Inventors")
    for item in [
        "US 10,123,456 — smart bottle with weight sensor only, no temperature.",
        "US 2019/0256123 A1 — temperature-monitoring mug, no volume sensing.",
        "Hidrate Spark 3 (commercial) — BLE hydration tracker, no temperature reading.",
    ]:
        b(doc, item)

    path = OUT / "test_id.docx"
    doc.save(str(path))
    print(f"Created: {path}")


# ── 2. Additional Information ─────────────────────────────────────────────────

def make_additional():
    doc = Document()
    h(doc, "Additional Information — Smart Bottle", 0)
    p(doc, "Date: 2024-01-12   |   Supplement to main Invention Disclosure")

    h(doc, "A. Prototype Bench-Test Results (Prototype v0.3, December 2023)")
    rows = [
        ("Metric",              "Achieved",  "Target"),
        ("Temperature accuracy", "±0.4 °C",  "±0.5 °C"),
        ("Volume accuracy",      "±1.8 ml",  "±2 ml"),
        ("BLE range (open air)", "42 m",      "≥ 30 m"),
        ("Battery life",         "16 days",   "≥ 14 days"),
        ("Water-resistance",     "IP67 pass", "IP67"),
    ]
    tbl = doc.add_table(rows=len(rows), cols=3)
    tbl.style = "Table Grid"
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            tbl.cell(r_idx, c_idx).text = cell_text

    h(doc, "B. Manufacturing Notes")
    p(doc, (
        "The TPU gasket is injection-moulded to a ±0.05 mm tolerance to achieve IP67. "
        "Each load cell is epoxy-bonded and individually calibrated on the production line "
        "using a two-point calibration (0 g and 500 g reference weights). "
        "The Nordic nRF52840 module is FCC/IC pre-certified, reducing regulatory lead time."
    ))

    h(doc, "C. Regulatory Considerations")
    for item in [
        "FCC Part 15 (BLE radio) — pre-certified module; declaration of conformity required.",
        "FDA food-contact materials compliance for inner stainless-steel surface coating.",
        "CE marking for EU distribution — EN 300 328 (BLE) + EN 62368-1 (safety); pending.",
        "RoHS / WEEE compliance confirmed for all PCB components.",
    ]:
        b(doc, item)

    h(doc, "D. Licensing Intent")
    p(doc, (
        "The inventors intend to license core sensing + firmware IP to established drinkware "
        "brands. The UI/UX design of the mobile app may be separately licensed or kept "
        "proprietary. Defensive publication is acceptable for peripheral features "
        "(e.g., RGB LED hydration indicator)."
    ))

    path = OUT / "test_additional.docx"
    doc.save(str(path))
    print(f"Created: {path}")


# ── 3. Patent Claims — with embedded Word comments (= verification questions) ─

_COMMENTS_XML = """\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments
    xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">

  <w:comment w:id="1" w:author="Patent Counsel" w:date="2024-01-20T09:00:00Z" w:initials="PC">
    <w:p>
      <w:r>
        <w:t>Does Claim 1 sufficiently define the positional relationship of the temperature sensing element to the liquid? Should the claim specify that the sensor is in direct thermal contact with the contents to support enablement under 35 U.S.C. § 112?</w:t>
      </w:r>
    </w:p>
  </w:comment>

  <w:comment w:id="2" w:author="Patent Counsel" w:date="2024-01-20T09:05:00Z" w:initials="PC">
    <w:p>
      <w:r>
        <w:t>Is the term 'wireless communication module' in Claim 3 broad enough to cover Wi-Fi and Zigbee implementations described in the disclosure, or should the claim language be updated to recite 'short-range wireless protocol' rather than BLE specifically?</w:t>
      </w:r>
    </w:p>
  </w:comment>

  <w:comment w:id="3" w:author="Patent Counsel" w:date="2024-01-20T09:10:00Z" w:initials="PC">
    <w:p>
      <w:r>
        <w:t>Claim 5 recites a 'hydration deficit calculation method' — does the disclosure provide sufficient algorithmic detail (e.g., the 35 ml/kg formula and threshold logic) to support this claim for written description and enablement purposes?</w:t>
      </w:r>
    </w:p>
  </w:comment>

</w:comments>
"""

_CT_OVERRIDE = (
    '<Override PartName="/word/comments.xml" '
    'ContentType="application/vnd.openxmlformats-officedocument'
    '.wordprocessingml.comments+xml"/>'
)

_REL_ENTRY = (
    '<Relationship Id="rIdComments" '
    'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" '
    'Target="comments.xml"/>'
)


def _inject_comments(docx_path: str):
    """Patch an existing .docx ZIP to add word/comments.xml and register it."""
    tmp = docx_path + ".tmp"
    with zipfile.ZipFile(docx_path, "r") as zin, \
         zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:

        for item in zin.infolist():
            data = zin.read(item.filename)

            if item.filename == "[Content_Types].xml":
                text = data.decode("utf-8")
                if "comments.xml" not in text:
                    text = text.replace("</Types>", _CT_OVERRIDE + "\n</Types>")
                data = text.encode("utf-8")

            elif item.filename == "word/_rels/document.xml.rels":
                text = data.decode("utf-8")
                if "comments" not in text.lower():
                    text = text.replace("</Relationships>", _REL_ENTRY + "\n</Relationships>")
                data = text.encode("utf-8")

            zout.writestr(item, data)

        zout.writestr("word/comments.xml", _COMMENTS_XML.encode("utf-8"))

    os.replace(tmp, docx_path)


def make_claim():
    doc = Document()
    h(doc, "Patent Claims — Smart Hydration Monitoring Bottle", 0)
    p(doc, "Application No.: [PENDING]   |   Filing Date: 2024-01-25")

    h(doc, "Independent Claims")

    h(doc, "Claim 1 — Apparatus", 2)
    p(doc, (
        "A portable hydration vessel comprising:\n"
        "  a vessel body defining an interior volume configured to hold a liquid;\n"
        "  a temperature sensing element configured to measure a temperature of the liquid;\n"
        "  a volume sensing element configured to measure a quantity of liquid in the vessel body;\n"
        "  a processor operatively coupled to both sensing elements; and\n"
        "  a non-transitory memory storing instructions that, when executed by the processor, "
        "cause the processor to generate hydration data based on the measured temperature "
        "and the measured quantity of liquid."
    ))

    h(doc, "Claim 3 — Apparatus with Wireless Communication", 2)
    p(doc, (
        "The portable hydration vessel of Claim 1, further comprising:\n"
        "  a wireless communication module configured to transmit the hydration data "
        "to an external computing device;\n"
        "  wherein the wireless communication module operates according to a "
        "short-range wireless protocol."
    ))

    h(doc, "Claim 5 — Method", 2)
    p(doc, (
        "A computer-implemented method for hydration monitoring, comprising:\n"
        "  receiving, from a sensor module of a portable vessel, temperature data and volume data;\n"
        "  computing a hydration deficit based on the volume data and a user-defined "
        "daily hydration target; and\n"
        "  generating a notification when the hydration deficit exceeds a predetermined threshold."
    ))

    h(doc, "Dependent Claims")

    h(doc, "Claim 2", 2)
    p(doc, (
        "The portable hydration vessel of Claim 1, wherein the temperature sensing element "
        "comprises a negative temperature coefficient (NTC) thermistor."
    ))

    h(doc, "Claim 4", 2)
    p(doc, (
        "The portable hydration vessel of Claim 1, wherein the volume sensing element "
        "comprises a strain-gauge load cell configured to measure a weight of the liquid."
    ))

    h(doc, "Claim 6", 2)
    p(doc, (
        "The method of Claim 5, wherein the daily hydration target is computed as "
        "35 millilitres per kilogram of a user's body weight."
    ))

    path = str(OUT / "test_claim.docx")
    doc.save(path)
    _inject_comments(path)
    print(f"Created: {path}  (3 Word comments embedded as questions)")


# ── entry point ───────────────────────────────────────────────────────────────

if __name__ == "__main__":
    make_id()
    make_additional()
    make_claim()
    print("\nAll three test documents created.")
