import zipfile
import xml.etree.ElementTree as ET
from docx import Document

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NS = {"w": _W_NS}


def extract_text(docx_path: str) -> str:
    """Extract all plain text from a .docx file (paragraphs + table cells)."""
    doc = Document(docx_path)
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(text)

    return "\n".join(parts)


def extract_comments(docx_path: str) -> list[str]:
    """Extract all comment texts from a .docx file (used as verification questions)."""
    comments: list[str] = []

    try:
        with zipfile.ZipFile(docx_path) as z:
            if "word/comments.xml" not in z.namelist():
                return comments

            with z.open("word/comments.xml") as f:
                root = ET.parse(f).getroot()

            for comment_el in root.findall(".//w:comment", _NS):
                text_nodes = comment_el.findall(".//w:t", _NS)
                text = " ".join(t.text for t in text_nodes if t.text).strip()
                if text:
                    comments.append(text)

    except Exception as exc:
        print(f"[document_processor] Comment extraction failed: {exc}")

    return comments
