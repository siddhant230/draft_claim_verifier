from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def _add_markdown_paragraph(doc: Document, line: str) -> None:
    """Add a single line to the document with basic markdown-to-docx conversion."""
    stripped = line.strip()

    if not stripped:
        doc.add_paragraph()
        return

    # Heading levels
    if stripped.startswith("### "):
        doc.add_heading(stripped[4:], level=3)
        return
    if stripped.startswith("## "):
        doc.add_heading(stripped[3:], level=2)
        return
    if stripped.startswith("# "):
        doc.add_heading(stripped[2:], level=1)
        return

    # Bullet list
    if stripped.startswith(("- ", "* ", "• ")):
        doc.add_paragraph(stripped[2:], style="List Bullet")
        return

    # Numbered list (e.g. "1. ")
    if len(stripped) > 2 and stripped[0].isdigit() and stripped[1] == "." and stripped[2] == " ":
        doc.add_paragraph(stripped[3:], style="List Number")
        return

    # Plain paragraph (handle inline **bold**)
    para = doc.add_paragraph()
    _add_inline_bold(para, stripped)


def _add_inline_bold(para, text: str) -> None:
    """Split text on **…** markers and add runs with bold toggled."""
    parts = text.split("**")
    for i, part in enumerate(parts):
        run = para.add_run(part)
        run.bold = (i % 2 == 1)


def save_analysis_to_docx(analysis_text: str, output_path: str) -> str:
    """Save the streaming analysis text to a formatted .docx file."""
    doc = Document()

    title = doc.add_heading("Patent Claim Analysis Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    for line in analysis_text.splitlines():
        _add_markdown_paragraph(doc, line)

    doc.save(output_path)
    return output_path


def save_qa_to_docx(qa_pairs: list[tuple[str, str]], output_path: str) -> str:
    """Save approved Q&A pairs to a formatted .docx file."""
    doc = Document()

    title = doc.add_heading("Patent Claim Verification — Q&A Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  |  "
        f"Total approved pairs: {len(qa_pairs)}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for i, (question, answer) in enumerate(qa_pairs, 1):
        # Question heading
        doc.add_heading(f"Question {i}", level=2)

        q_para = doc.add_paragraph()
        q_run = q_para.add_run("Q: ")
        q_run.bold = True
        q_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        q_para.add_run(question)

        doc.add_paragraph()

        a_heading = doc.add_paragraph()
        a_run = a_heading.add_run("Answer:")
        a_run.bold = True

        for line in answer.splitlines():
            _add_markdown_paragraph(doc, line)

        # Divider
        doc.add_paragraph()
        doc.add_paragraph("─" * 60)
        doc.add_paragraph()

    doc.save(output_path)
    return output_path
